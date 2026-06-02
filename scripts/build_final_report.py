from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "final-report"
OUT_PATH = OUT_DIR / "Sport_Analysis_Dashboard_최종보고서.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr_cells[i], "1E293B")
        if widths:
            hdr_cells[i].width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            set_cell_shading(cells[i], "F8FAFC" if len(table.rows) % 2 else "FFFFFF")
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 18, "0F172A"),
        ("Heading 2", 14, "1D4ED8"),
        ("Heading 3", 11, "334155"),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sport Analysis Dashboard")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("웹프로그래밍2 팀 프로젝트 최종 보고서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = summary.add_run(
        "축구, 야구, e스포츠 데이터를 통합해 경기 일정, 순위, 상세 기록, 팬 참여, AI 경기 분석을 제공하는 스포츠 분석 웹 서비스"
    )
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(11)

    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "Sport Analysis Dashboard"],
            ["기술 스택", "Spring Boot, React, Oracle DB, Vite, Groq API, MLB Stats API, Cito API"],
            ["주요 사용자", "스포츠 팬, 경기 기록 확인 사용자, 관리자"],
            ["보고서 작성일", "2026-06-02"],
            ["주의", "개인별 소감 및 성찰은 각 팀원이 직접 작성해야 합니다."],
        ],
        [Cm(4), Cm(12)],
    )
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        run = p.add_run(item)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Number"]
        run = p.add_run(item)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF2FF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(67, 56, 202)
    r2 = p.add_run(body)
    r2.font.name = "Malgun Gothic"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("1. 프로젝트 개요", level=1)
    doc.add_paragraph(
        "Sport Analysis Dashboard는 축구(K리그), 야구(MLB), e스포츠(LCK)를 하나의 기준으로 탐색할 수 있도록 만든 통합 스포츠 분석 웹 서비스이다. "
        "사용자는 경기 일정과 결과를 확인하고, 상세 페이지에서 라인스코어·선수 지표·팬 투표·채팅·AI 분석을 이어서 볼 수 있다."
    )
    add_table(
        doc,
        ["구분", "목표", "구현 결과"],
        [
            ["정보 통합", "종목별로 흩어진 경기 데이터를 한 서비스에서 확인", "경기센터, 종목 허브, 순위 화면으로 통합"],
            ["상세 분석", "경기 종료 후 주요 기록과 분석 제공", "MLB 상세 기록, LCK 경기/게임별 통계, K리그 분석"],
            ["사용자 참여", "즐겨찾기, 팬 투표, 채팅으로 참여성 강화", "로그인 기반 즐겨찾기, 팬 승부 예측, 경기 채팅방"],
            ["관리", "관리자가 외부 데이터를 동기화", "축구, MLB, e스포츠 동기화 버튼과 권한 검증"],
        ],
        [Cm(3), Cm(6.5), Cm(6.5)],
    )

    doc.add_heading("2. 평가 항목 대응 요약", level=1)
    add_table(
        doc,
        ["평가 항목", "대응 내용", "증거 파일"],
        [
            ["산출물 완성도", "React SPA와 Spring Boot REST API 기반 기능 구현", "frontend/src/pages, src/main/java"],
            ["시스템 설계", "종목별 모듈 분리와 공통 Match/Team 구조 병행", "docs/MLB_DESIGN_DOCS.md, README-soccer.md"],
            ["AI 활용 투명성", "프롬프트 로그와 AI 기여도 표로 사용 범위 명시", "docs/*PROMPT_LOG.md"],
            ["검증 및 개선", "프론트 빌드, 백엔드 컴파일, 화면별 체크리스트 관리", "docs/FINAL_TEST_CHECKLIST.md"],
            ["발표 가능성", "시연 순서와 장애 대응 포인트 정리", "본 보고서 10장, 발표 자료"],
        ],
        [Cm(3.2), Cm(7.2), Cm(5.6)],
    )

    doc.add_heading("3. 기술 스택 및 아키텍처", level=1)
    add_table(
        doc,
        ["영역", "기술", "선택 이유"],
        [
            ["Frontend", "React, Vite, Axios", "화면 상태 관리와 API 연동을 분리하기 쉽고 개발 서버 반응이 빠름"],
            ["Backend", "Spring Boot, JPA", "REST API 서버와 Oracle DB 연동에 적합"],
            ["Database", "Oracle DB", "수업 환경과 기존 프로젝트 구조를 유지"],
            ["AI", "Groq API, Gemini 보조 구조", "경기 결과 요약과 전술 분석 생성"],
            ["외부 데이터", "MLB Stats API, Cito API, TheSportsDB", "종목별 실제 경기 데이터 확보"],
        ],
        [Cm(3), Cm(5), Cm(8)],
    )
    add_callout(
        doc,
        "핵심 설계 판단",
        "종목별 API 응답 구조가 크게 다르기 때문에 축구, 야구, e스포츠 패키지를 분리했다. 다만 사용자 화면에서는 공통 경기센터와 순위 화면으로 연결해 서비스 경험은 하나로 유지했다.",
    )

    doc.add_heading("4. 주요 기능", level=1)
    add_table(
        doc,
        ["화면/기능", "설명"],
        [
            ["홈", "실제 LIVE 상태의 경기만 실시간으로 표시하고 3일 내 예정 경기를 요약"],
            ["경기센터", "날짜, 종목, 상태 필터를 조합해 경기 목록 조회"],
            ["경기 상세", "스코어보드, 라인스코어, 핵심 타자, 선발 투수, 채팅, 팬 투표 제공"],
            ["AI 분석", "종료 경기 중심으로 분석 가능한 경기 목록과 결과 제공"],
            ["스포츠 허브", "축구, 야구, e스포츠 전용 화면으로 진입"],
            ["관리자 대시보드", "외부 데이터 동기화와 DB 반영 여부 확인"],
            ["설정/내 팀", "프로필, 비밀번호, 즐겨찾기 팀 관리"],
        ],
        [Cm(4), Cm(12)],
    )

    doc.add_heading("5. 종목별 구현 내용", level=1)
    add_table(
        doc,
        ["종목", "구현 내용", "차별점"],
        [
            ["축구", "K리그 일정, 순위, 상세 경기, AI 분석", "K리그 실제 데이터 기반 화면과 시즌별 순위"],
            ["야구", "MLB 일정, 상세 라인스코어, 투구/타격 기록, AI 리포트", "MLB Stats API 기반 세부 기록 시각화"],
            ["e스포츠", "LCK 일정, 시즌 결과, 팀 정보, 게임별 통계, Groq 분석", "Cito API와 DB 선수 지표를 결합한 분석"],
        ],
        [Cm(2.5), Cm(7), Cm(6.5)],
    )

    doc.add_heading("6. AI 활용 방식과 학술적 투명성", level=1)
    doc.add_paragraph(
        "AI는 코드 초안 작성, 오류 원인 탐색, UI 개선 방향 제안, 문서 초안 정리에 사용했다. "
        "팀원은 요구사항 확정, 실행 결과 확인, 오류 재현, API 키 관리, 최종 채택 여부 판단을 담당했다."
    )
    add_table(
        doc,
        ["단계", "AI 활용", "사람의 검토"],
        [
            ["기획", "기능 범위와 구조 제안", "팀 일정과 평가 기준에 맞춰 우선순위 조정"],
            ["구현", "API, 컴포넌트, DTO 초안 작성", "실제 DB 컬럼과 응답 형식에 맞게 수정"],
            ["검증", "체크리스트와 오류 수정안 제시", "브라우저 화면과 빌드 결과로 직접 확인"],
            ["문서", "보고서 초안과 발표 흐름 정리", "개인 성찰과 사실관계는 팀원이 직접 보완"],
        ],
        [Cm(3), Cm(6.5), Cm(6.5)],
    )
    add_callout(
        doc,
        "보고서 작성 주의",
        "개인별 느낀 점, 학습 경험, 역할 회고는 AI가 대신 작성하면 안 되므로 제출 전 각 팀원이 직접 자신의 표현으로 보완해야 한다.",
    )

    doc.add_heading("7. 주요 문제 해결 사례", level=1)
    add_numbered(
        doc,
        [
            "관리자 대시보드에서 존재하지 않는 필드를 map으로 읽어 빈 화면이 발생하던 문제를 방어 코드로 수정했다.",
            "Maven wrapper가 Windows에서 .m2 경로 Target[0]을 읽다가 실패하던 문제를 수정해 백엔드 컴파일을 통과시켰다.",
            "실시간 경기 카운트가 실제 LIVE가 아닌 경기를 포함하던 문제를 실제 상태 기준으로 정리했다.",
            "팀 로고 표시 방식이 화면마다 달라 보이던 문제를 공통 TeamLogo 컴포넌트 기준으로 통일했다.",
            "e스포츠 오늘 경기에서 TBD placeholder가 보이던 문제를 제외 조건으로 처리했다.",
        ]
    )

    doc.add_heading("8. 검증 결과", level=1)
    add_table(
        doc,
        ["검증 항목", "명령/방법", "결과"],
        [
            ["백엔드 컴파일", ".\\mvnw.cmd -q -DskipTests compile", "PASS"],
            ["프론트 린트", "cd frontend && npm.cmd run lint", "PASS"],
            ["프론트 빌드", "cd frontend && npm.cmd run build", "PASS"],
            ["화면 검증", "홈, 경기센터, 스포츠 허브, 분석, 상세 화면 확인", "주요 오류 정리"],
        ],
        [Cm(4), Cm(8), Cm(4)],
    )

    doc.add_heading("9. 현실적 제한과 향후 개선", level=1)
    add_bullets(
        doc,
        [
            "외부 API 키가 없거나 호출 제한에 걸리면 최신 데이터 동기화가 실패할 수 있다.",
            "Oracle DB 로컬 설정이 필요하므로 발표 환경에서는 DB 접속 정보를 사전에 점검해야 한다.",
            "AI 분석은 경기 결과 요약 보조 기능이며, 승부 예측이나 사실 확정 자료로 오해하지 않도록 안내가 필요하다.",
            "향후에는 통합 테스트와 배포 환경, 모바일 반응형 최적화를 추가하면 완성도를 더 높일 수 있다.",
        ]
    )

    doc.add_heading("10. 발표 시연 흐름 제안", level=1)
    add_numbered(
        doc,
        [
            "홈에서 실시간 경기와 3일 내 예정 경기 요약을 보여준다.",
            "경기센터에서 날짜, 종목, 상태 필터가 실제 목록을 바꾸는 모습을 보여준다.",
            "MLB 상세 페이지에서 라인스코어, 핵심 타자, 팬 투표, 채팅 UI를 시연한다.",
            "e스포츠 화면에서 LCK 시즌, 경기 결과, 팀 정보, AI 분석 흐름을 보여준다.",
            "관리자 대시보드에서 데이터 동기화 버튼과 권한 제한을 설명한다.",
            "마지막으로 AI 사용 투명성과 검증 결과를 제시한다.",
        ]
    )

    doc.add_heading("부록. 제출 전 확인 목록", level=1)
    add_bullets(
        doc,
        [
            "팀원 이름, 학번, 역할 분담 표를 실제 제출 형식에 맞게 추가한다.",
            "개인 성찰 문단은 각 팀원이 직접 작성한다.",
            "발표 당일에는 백엔드, 프론트엔드, DB, API 키를 먼저 실행 확인한다.",
            "GitHub 제출 전 node_modules, dist, target, 개인 API 키가 포함되지 않았는지 확인한다.",
        ]
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
