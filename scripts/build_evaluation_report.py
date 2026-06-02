from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "final-report"
OUT_PATH = OUT_DIR / "Sport_Analysis_Dashboard_평가기준_최종보고서.docx"


def east_asia(run, font="Malgun Gothic"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="CBD5E1"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        tc_borders.append(tag)


def set_cell(cell, text, bold=False, color="0F172A", align=None, fill=None):
    cell.text = ""
    if fill:
        shade(cell, fill)
    borders(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align or (WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 16 else WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(str(text))
    east_asia(r)
    r.bold = bold
    r.font.size = Pt(8.6)
    r.font.color.rgb = RGBColor.from_string(color)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, color="FFFFFF", fill="1E293B")
        if widths:
            t.rows[0].cells[i].width = widths[i]
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        fill = "F8FAFC" if ridx % 2 == 0 else "FFFFFF"
        for i, v in enumerate(row):
            set_cell(cells[i], v, fill=fill)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return t


def paragraph(doc, text="", bold=False, size=10, color="111827", align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    east_asia(r)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    east_asia(r)
    r.font.size = Pt(9.5)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    east_asia(r)
    r.font.size = Pt(9.5)


def callout(doc, title, body, fill="EEF2FF"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    borders(cell, "A5B4FC")
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    east_asia(r)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("4338CA")
    r2 = p.add_run(body)
    east_asia(r2)
    r2.font.size = Pt(9.2)
    r2.font.color.rgb = RGBColor.from_string("1E293B")
    doc.add_paragraph()


def section_title(doc, text, po=None):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        east_asia(run)
        run.font.color.rgb = RGBColor.from_string("0F172A")
        run.font.size = Pt(17)
    if po:
        paragraph(doc, po, bold=True, size=9, color="2563EB", space_after=8)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        east_asia(run)
        run.font.color.rgb = RGBColor.from_string("1D4ED8")
        run.font.size = Pt(13)


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        east_asia(run)
        run.font.color.rgb = RGBColor.from_string("334155")
        run.font.size = Pt(11)


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    paragraph(doc, "프로젝트 결과 보고서", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    paragraph(doc, "Sport Analysis Dashboard", bold=True, size=20, color="1D4ED8", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    paragraph(
        doc,
        "축구, 야구, e스포츠 경기 데이터를 통합하고 경기 상세 기록, 팬 참여, AI 분석을 제공하는 스포츠 분석 웹 서비스",
        size=11,
        color="475569",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )
    table(
        doc,
        ["항목", "내용"],
        [
            ["교과목명", "웹 프로그래밍 2"],
            ["프로젝트명", "Sport Analysis Dashboard"],
            ["팀명", "Sport Analysis Team / 실제 팀명 기입 필요"],
            ["팀원", "조우성, 김우림, 최수인 / 실제 학번·이름 확인 필요"],
            ["담당교수", "기입 필요"],
            ["제출일", "2026-06-02 / 실제 제출일 확인 필요"],
            ["학기", "2026학년도 1학기"],
        ],
        [Cm(4), Cm(12)],
    )
    callout(
        doc,
        "작성 원칙",
        "본 보고서는 제공된 팀프로젝트 보고서 양식과 웹프로그래밍2 팀프로젝트 평가기준을 기준으로 재작성하였다. "
        "개인 성찰 항목은 평가기준상 AI 작성이 금지되어 있으므로 초안이 아닌 작성 지침과 빈 칸으로 남겼다.",
        "F0F9FF",
    )
    h2(doc, "평가 루브릭 직접 대응표")
    table(
        doc,
        ["평가 영역", "평가 항목", "배점", "보고서 대응 위치"],
        [
            ["산출물", "① 시스템 완성도 및 기능 구현", "15%", "1장, 5장, 8장"],
            ["산출물", "② 아키텍처 및 설계 판단", "15%", "1.3, 4장"],
            ["과정", "③ AI 활용 전략 및 프롬프트 엔지니어링", "15%", "2장, 3장"],
            ["과정", "④ 코드 검증 및 비판적 사고", "10%", "5.2, 5.3, 8.2"],
            ["과정", "⑤ 팀워크 및 의사소통", "15%", "6장"],
            ["과정", "⑥ AI 사용 투명성 및 학술적 정직성", "5%", "2장, 7.2, 9장"],
            ["성찰", "⑦ 학습 성찰 보고서", "15%", "9장: 팀원 직접 작성"],
            ["성찰", "⑧ 발표 및 시연", "10%", "8장, 발표 PPT, 발표 대본"],
        ],
        [Cm(2.7), Cm(6.3), Cm(2), Cm(5)],
    )
    doc.add_page_break()

    section_title(doc, "1. 프로젝트 개요 및 문제 정의", "[PO3] 문제 정의 및 모델링, [PO5] 현실적 제한조건 고려 설계")
    h2(doc, "1.1 문제 정의 및 프로젝트 배경")
    paragraph(
        doc,
        "스포츠 팬은 경기 일정, 결과, 순위, 선수 기록을 확인하기 위해 종목별로 서로 다른 사이트와 데이터 형식을 오가야 한다. "
        "축구, 야구, e스포츠는 경기 단위, 기록 항목, 상태 표현 방식이 다르기 때문에 단순 목록 화면만으로는 통합 서비스 경험을 만들기 어렵다. "
        "또한 경기 결과가 나온 뒤에도 사용자는 단순 스코어를 넘어 어떤 선수가 중요한 역할을 했는지, 경기 흐름이 어떻게 바뀌었는지, 팀별 순위 기준이 일관적인지 알고 싶어 한다.",
    )
    paragraph(
        doc,
        "본 프로젝트는 이러한 문제를 해결하기 위해 종목별 실제 데이터를 하나의 웹 서비스에서 탐색하고, 경기 상세 기록과 AI 요약을 통해 결과를 이해할 수 있는 Sport Analysis Dashboard를 개발하였다.",
    )
    h2(doc, "1.2 프로젝트 목표 및 범위")
    table(
        doc,
        ["목표", "구현 범위", "평가기준 대응"],
        [
            ["통합 경기 탐색", "홈, 경기센터, 스포츠 허브에서 축구·야구·e스포츠 경기 조회", "시스템 완성도"],
            ["상세 기록 제공", "MLB 라인스코어·선발 투수·핵심 타자, LCK 게임별 통계, K리그 상세", "기능 구현"],
            ["AI 분석", "종료 경기 중심의 Groq/Gemini 기반 경기 요약·전술 분석", "최신 도구 활용"],
            ["사용자 참여", "즐겨찾기, 팬 승부 예측, 경기 채팅", "사용자 요구 반영"],
            ["관리 기능", "관리자 권한, 외부 데이터 동기화, API 상태 확인", "보안·운영 고려"],
        ],
        [Cm(4), Cm(7), Cm(4.5)],
    )
    h2(doc, "1.3 현실적 제한조건 분석")
    table(
        doc,
        ["제약", "내용", "대응"],
        [
            ["경제성", "상용 배포·유료 API 사용을 최소화해야 함", "무료/공개 API와 로컬 개발 환경 중심으로 구성"],
            ["기술적 제약", "Oracle DB, Spring Boot, React를 수업 및 기존 프로젝트 구조에 맞춰 사용", "기존 Entity/Repository 구조를 최대한 유지"],
            ["시간적 제약", "짧은 팀 프로젝트 기간 내 세 종목을 모두 연결해야 함", "종목별 담당 범위를 분리하고 공통 UI는 재사용"],
            ["안전성", "API 키, 로그인 세션, 관리자 권한 보호 필요", "환경변수 사용, HttpSession 인증, AdminRoute 적용"],
            ["AI 도구 제약", "AI가 생성한 코드가 실제 DB/API와 불일치할 수 있음", "빌드 검증, 화면 확인, 프롬프트 로그와 체크리스트 운영"],
            ["사회적 영향", "AI 분석이 실제 예측처럼 오해될 수 있음", "팬 투표와 AI 분석을 구분하고 보조 분석임을 명시"],
        ],
        [Cm(3), Cm(6.3), Cm(6.7)],
    )

    section_title(doc, "2. AI 도구 사용 투명성 보고", "[PO4] 최신 도구 활용, [PO9] 직업윤리와 사회적 책임")
    h2(doc, "2.1 사용 도구 명세")
    table(
        doc,
        ["도구명", "모델/버전", "사용 기간", "사용 범위", "비용"],
        [
            ["Claude Code / Codex", "대화형 코딩 에이전트", "2026년 5월~6월", "코드 작성, 버그 수정, 문서화, 발표자료 초안", "구독/플랫폼 사용"],
            ["Groq API", "Chat Completions", "구현 단계", "경기 결과 요약 및 전술 분석", "API 키 필요"],
            ["Gemini API", "보조 분석 구조", "축구 분석 실험", "K리그 분석 보조", "API 키 필요"],
            ["Canva", "발표 디자인 보조 가능", "제출물 단계", "PPT 구성을 Canva로 옮길 수 있도록 구조화", "계정/서비스 조건에 따름"],
        ],
        [Cm(3.2), Cm(3), Cm(2.8), Cm(5.5), Cm(2.4)],
    )
    h2(doc, "2.2 AI 기여도 명세")
    table(
        doc,
        ["모듈/기능", "AI 기여", "팀원 기여", "검토/채택 기준"],
        [
            ["React 전환", "라우터·컴포넌트 구조 초안", "기존 기능 유지 여부 판단", "빌드와 화면 확인 후 채택"],
            ["인증/권한", "Controller, DTO, Route Guard 구현 초안", "HttpSession 유지 결정", "Spring Security 미도입 원칙 반영"],
            ["경기센터", "필터 UI와 API 연동 코드 수정", "날짜/상태 기준 오류 재현", "실제 LIVE/예정/종료 기준 검증"],
            ["MLB 상세", "라인스코어·선발·타자·AI 탭 컴포넌트 작성", "UI 시인성 문제 발견", "MLB Stats API 응답과 대조"],
            ["LCK/e스포츠", "Cito API 프록시와 화면 구성 보조", "기존 브랜치 기능 비교 요구", "TBD 제외, 오늘 경기 기준 확인"],
            ["문서화", "보고서·PPT 초안 생성", "팀원별 사실관계·성찰 보완 필요", "평가기준 문서와 대조"],
        ],
        [Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.5)],
    )
    h2(doc, "2.3 세션 로그 요약")
    paragraph(doc, "핵심 프롬프트는 전체 로그 파일에 보관했으며, 보고서에는 평가 기준과 연결되는 대표 사례를 발췌했다.")
    table(
        doc,
        ["사례", "프롬프트 목적", "AI 응답 평가", "후속 조치"],
        [
            ["React REST 전환", "Spring Boot를 REST API 서버로 정리하고 React 화면 연결", "채택", "라우터, Axios API, CORS, VITE_API_BASE_URL 정리"],
            ["관리자 오류 수정", "undefined map 오류로 /admin 빈 화면 발생", "채택", "방어 코드와 ErrorBoundary 추가"],
            ["경기센터 기준 수정", "LIVE/예정/종료 필터와 날짜 버튼 오류 개선", "부분 채택", "실제 데이터 기준을 다시 확인하며 여러 차례 수정"],
            ["Maven wrapper 복구", "mvnw.cmd 실패 원인 분석", "채택", "Windows PowerShell Target[0] 오류 수정 후 compile PASS"],
            ["보고서/PPT 작성", "평가기준 기반 산출물 작성", "채택", "개인 성찰은 직접 작성하도록 분리"],
        ],
        [Cm(3), Cm(5.5), Cm(2.5), Cm(6)],
    )

    section_title(doc, "3. 프롬프트 엔지니어링 및 AI 협업 전략", "[PO4] 최신 도구 활용, [PO7] 의사소통")
    h2(doc, "3.1 프롬프트 전략의 진화")
    table(
        doc,
        ["구분", "프로젝트 초기", "프로젝트 후기"],
        [
            ["구체성", "기능을 만들어 달라는 단순 요청 중심", "목표, 파일, 제약조건, 검증 명령을 함께 제시"],
            ["맥락 제공", "현재 DB/API 구조 설명이 부족", "기존 Entity, 외부 API 응답, 화면 오류 캡처를 제공"],
            ["제한조건", "AI가 새 라이브러리나 구조 변경을 제안할 위험", "Spring Security 미도입, 기존 구조 유지, git commit 금지 등 반복 명시"],
            ["검증 요청", "구현 후 확인 기준이 모호", "npm lint/build, mvnw compile, 브라우저 스모크 테스트를 요구"],
            ["성장 증거", "결과물 중심", "오류 재현 → 원인 분리 → 수정 → 검증의 루프 기록"],
        ],
        [Cm(3), Cm(6), Cm(7)],
    )
    h2(doc, "3.2 설계 단계별 프롬프트 기록")
    h3(doc, "3.2.1 목표/기준 설정 단계")
    paragraph(doc, "프롬프트 요지: React + Spring Boot REST API 구조로 전환하되 기존 Entity, Repository, Service 구조를 최대한 유지하고 Spring Security는 추가하지 않는다.")
    paragraph(doc, "평가: 채택. 초기 구조 전환 비용은 있었지만 이후 프론트 화면 개선과 API 검증이 쉬워졌다.")
    h3(doc, "3.2.2 분석/설계 단계")
    paragraph(doc, "프롬프트 요지: 관리자 권한, ProtectedRoute, AdminRoute, 세션 기반 권한 확인을 설계한다.")
    paragraph(doc, "평가: 채택. JWT보다 확장성은 낮지만 팀 프로젝트 기간과 기존 세션 구조를 고려하면 현실적인 선택이었다.")
    h3(doc, "3.2.3 제작/구현 단계")
    paragraph(doc, "프롬프트 요지: MLB Stats API, Cito LCK API, K리그 데이터를 각각 동기화하고 공통 화면에서 표시한다.")
    paragraph(doc, "평가: 부분 채택. AI가 제안한 코드 중 실제 API 응답과 다른 부분은 팀원이 재현 화면과 오류 메시지를 제공해 수정했다.")
    h3(doc, "3.2.4 시험/테스트 단계")
    paragraph(doc, "프롬프트 요지: 관리자 대시보드 undefined 오류, 경기센터 날짜 필터 오류, 팀 로고 누락, Maven wrapper 오류를 재현 조건과 함께 수정 요청한다.")
    paragraph(doc, "평가: 채택. 단순 수정이 아니라 오류 원인을 분리하고 검증 명령으로 확인했다.")
    h3(doc, "3.2.5 평가/개선 단계")
    paragraph(doc, "프롬프트 요지: UI 시인성, 문구 자연스러움, 발표 가능성, 보고서/PPT 구조를 개선한다.")
    paragraph(doc, "평가: 채택. 발표용으로 기능 나열보다 설계 판단과 검증 증거를 중심으로 재구성했다.")
    h2(doc, "3.3 프롬프트 전후 비교 분석")
    table(
        doc,
        ["평가 기준", "초기 방식", "개선된 방식"],
        [
            ["방어적 프롬프팅", "없음", "기존 기능 제거 금지, 실제 LIVE 기준, 예정 경기 3일 기준 등 반복"],
            ["컨텍스트 문서", "즉석 대화 중심", "README, PROMPT_LOG, FINAL_TEST_CHECKLIST, 평가 로그 활용"],
            ["검증 루프", "AI 응답 수용 후 화면 확인", "명령 실행, 브라우저 확인, 스크린샷 오류 재현"],
        ],
        [Cm(4), Cm(5), Cm(7)],
    )
    h2(doc, "3.4 워크플로 설계")
    for item in [
        "요구사항 정의: 사용자가 화면 캡처와 함께 문제를 제시",
        "원인 분석: 데이터 문제, UI 문제, API 문제를 분리",
        "구현: 최소 범위로 코드 수정 후 기존 기능 유지 확인",
        "검증: lint/build/compile과 주요 화면 스모크 테스트 수행",
        "기록: README, 체크리스트, 프롬프트 로그, 보고서에 반영",
    ]:
        numbered(doc, item)

    section_title(doc, "4. 시스템 설계 및 아키텍처", "[PO5] 사용자 요구사항과 현실적 제한조건을 고려한 설계")
    h2(doc, "4.1 시스템 아키텍처 설계")
    table(
        doc,
        ["계층", "구성", "역할"],
        [
            ["Frontend", "React, Vite, React Router, Axios", "화면 렌더링, 필터 상태, API 호출, 로그인 상태별 접근 제어"],
            ["Backend", "Spring Boot, Controller, Service, Repository", "REST API 제공, 외부 API 프록시, DB 저장/조회"],
            ["Database", "Oracle DB, JPA Entity", "사용자, 팀, 경기, 순위, 분석, 채팅, 투표 데이터 저장"],
            ["External API", "MLB Stats API, Cito API, TheSportsDB/K리그", "실제 경기 일정, 결과, 순위, 세부 기록 수집"],
            ["AI API", "Groq, Gemini 보조", "경기 결과 요약, 전술 분석, 핵심 포인트 생성"],
        ],
        [Cm(3), Cm(5.5), Cm(7.5)],
    )
    h2(doc, "4.2 설계 결정 기록")
    table(
        doc,
        ["결정", "채택 이유", "트레이드오프"],
        [
            ["React SPA + REST API", "종목별 화면 확장과 API 검증이 쉬움", "초기 전환 비용 증가"],
            ["HttpSession 인증", "기존 구조와 팀 프로젝트 기간에 적합", "JWT보다 외부 클라이언트 확장성 낮음"],
            ["종목별 패키지 분리", "API 응답 구조가 매우 달라 독립 구현이 안전", "일부 DTO/화면 로직 중복 발생"],
            ["AI 분석은 요약 중심", "승부 예측처럼 오해될 위험을 줄임", "예측 기능은 향후 과제로 분리"],
            ["VITE_API_BASE_URL", "개발 포트와 API 주소 전환이 쉬움", "환경변수 설정 누락 시 혼선 가능"],
        ],
        [Cm(4), Cm(6), Cm(6)],
    )
    h2(doc, "4.3 설계 과정에서 AI의 한계")
    for item in [
        "AI는 실제 Oracle DB 접속 상태, API 키 권한, 로컬 포트 충돌을 직접 알 수 없어 사용자의 실행 결과가 필요했다.",
        "화면 시인성 문제는 코드만 보고 판단하기 어려워 스크린샷 기반 반복 수정이 필요했다.",
        "기존 브랜치에 있던 e스포츠 기능과 현재 코드의 차이는 저장소 상태를 직접 비교해야 하므로 AI의 기억만으로는 불충분했다.",
        "AI가 제안한 방식이 실제 외부 API 응답과 맞지 않는 경우가 있어, 응답 필드 확인과 방어 코드가 필요했다.",
    ]:
        bullet(doc, item)

    section_title(doc, "5. 구현 및 코드 검증", "[PO1] 컴퓨팅 지식 응용, [PO2] 프로그래밍 검증")
    h2(doc, "5.1 구현 과정")
    table(
        doc,
        ["기능", "구현 내용", "검증/개선"],
        [
            ["홈", "실제 LIVE 경기와 3일 내 예정 경기 요약", "LIVE가 아닌 경기 표시 문제 수정"],
            ["경기센터", "날짜·종목·상태 필터, 상세 이동", "날짜 버튼/필터 기준 반복 수정"],
            ["야구 상세", "라인스코어, 선발 투수, 핵심 타자, 팬 투표, 채팅", "시인성 개선 및 불필요한 LIVE 문구 제거"],
            ["e스포츠", "LCK 오늘 경기, 시즌 결과, 팀 정보, AI 분석", "TBD 제외, Cito 오늘 날짜 기준 수정"],
            ["순위", "축구·야구·e스포츠 순위 통합", "스포츠 탭과 순위 탭 기준 일치"],
            ["관리자", "데이터 동기화, API 상태, 권한 제어", "undefined 필드 방어와 오류 표시 개선"],
            ["설정/인증", "로그인, 회원가입, 비밀번호 변경, 회원 탈퇴 UI", "미구현 기능 정리와 실제 기능 연결"],
        ],
        [Cm(3), Cm(7), Cm(6)],
    )
    h2(doc, "5.2 AI 생성 코드 리뷰 프로세스")
    for item in [
        "AI가 생성한 코드를 바로 채택하지 않고, 기존 파일 구조와 import 경로가 맞는지 확인했다.",
        "사용자가 브라우저에서 발견한 오류를 스크린샷과 URL로 재현해 수정 범위를 좁혔다.",
        "AI가 제안한 기능 중 실제 구현되지 않은 테마 색상, 더미 알림, 불필요한 종료 버튼 등은 제거하거나 문구를 정리했다.",
        "팀 로고, 경기 상태, 분석 가능 수치처럼 여러 화면에서 기준이 달라지는 항목은 공통 컴포넌트와 공통 API 기준으로 정리했다.",
    ]:
        bullet(doc, item)
    h2(doc, "5.3 테스팅 전략")
    table(
        doc,
        ["검증 항목", "방법", "결과"],
        [
            ["프론트 린트", "npm.cmd run lint", "PASS"],
            ["프론트 빌드", "npm.cmd run build", "PASS"],
            ["백엔드 컴파일", ".\\mvnw.cmd -q -DskipTests compile", "PASS"],
            ["문서 검증", "DOCX 내부 텍스트/표 확인, PPTX 슬라이드 개수 확인", "PASS"],
            ["렌더 검증", "LibreOffice 기반 DOCX 렌더 시도", "PC에 soffice 없음으로 미수행"],
        ],
        [Cm(4), Cm(7), Cm(5)],
    )
    callout(
        doc,
        "AI 오류 발견 사례",
        "백엔드 코드 오류가 아니라 Maven wrapper 스크립트의 Windows PowerShell 처리 문제로 compile 검증이 막혔다. "
        "mvnw.cmd에서 .m2 경로의 Target[0]을 무조건 접근하던 부분을 null/empty 방어 처리하여 backend compile을 통과시켰다.",
        "FFF7ED",
    )

    section_title(doc, "6. 팀 협업 및 역할 분담", "[PO6] 팀워크, [PO7] 의사소통")
    h2(doc, "6.1 팀원별 역할 및 기여도")
    table(
        doc,
        ["이름", "역할", "담당 작업", "AI 활용 방식", "기여도"],
        [
            ["조우성", "축구 담당", "K리그 데이터, 축구 상세, 축구 순위, 축구 문서", "K리그 설계/분석 보조", "기입/확인 필요"],
            ["김우림", "e스포츠 담당", "LCK 일정, Cito API, 팀/선수 정보, e스포츠 분석", "Cito 프록시·UI 보조", "기입/확인 필요"],
            ["최수인", "야구 담당", "MLB 일정, 상세 기록, 라인스코어, AI 리포트", "MLB API·컴포넌트 보조", "기입/확인 필요"],
            ["공통", "통합", "인증, 경기센터, 관리자, 즐겨찾기, 문서, 발표", "버그 수정·문서화 보조", "팀 협의 필요"],
        ],
        [Cm(2.3), Cm(2.5), Cm(5.1), Cm(5), Cm(2.2)],
    )
    h2(doc, "6.2 팀 협업 및 의사소통")
    for item in [
        "종목별 담당 범위를 분리하고 공통 화면에서 충돌하는 데이터 기준을 통합했다.",
        "AI 사용 시 기존 기능을 임의로 제거하지 않는다는 원칙을 반복 적용했다.",
        "화면 오류는 URL, 스크린샷, 기대 동작을 함께 공유하여 재현 가능하게 만들었다.",
        "보고서의 개인 성찰과 실제 기여도는 각 팀원이 직접 사실관계에 맞게 보완해야 한다.",
    ]:
        bullet(doc, item)
    h2(doc, "6.3 프로젝트 관리")
    table(
        doc,
        ["단계", "주요 작업", "산출물"],
        [
            ["기획", "문제 정의, 기능 범위, 종목별 담당 분리", "계획서, 평가 대응 로그"],
            ["설계", "REST API 전환, DB/API 구조, 화면 라우팅 설계", "REACT_REST_MIGRATION_PLAN, 설계 문서"],
            ["구현", "축구·야구·e스포츠 및 공통 기능 개발", "src/main/java, frontend/src"],
            ["검증", "빌드, 컴파일, 브라우저 화면 확인", "FINAL_TEST_CHECKLIST"],
            ["제출", "README, 보고서, PPT, 대본, zip 정리", "outputs/final-report"],
        ],
        [Cm(3), Cm(7), Cm(6)],
    )

    section_title(doc, "7. 사회적 영향 및 직업윤리", "[PO8] 사회적 영향, [PO9] 직업윤리")
    h2(doc, "7.1 프로젝트 결과물의 사회적 영향")
    for item in [
        "긍정적 영향: 여러 종목의 경기 정보를 한 화면에서 비교할 수 있어 스포츠 팬의 정보 접근성을 높인다.",
        "부정적 가능성: AI 분석이 실제 승부 예측이나 공식 통계처럼 오해될 수 있다.",
        "접근성: 색상 대비와 팀 로고 통일을 개선했지만 모바일 접근성 검증은 추가 과제로 남았다.",
        "경제성: 무료 API와 로컬 실행 중심 구조를 사용해 팀 프로젝트 비용 부담을 줄였다.",
    ]:
        bullet(doc, item)
    h2(doc, "7.2 AI 도구 사용의 윤리적 고려")
    for item in [
        "API 키와 개인정보를 AI에게 직접 입력하지 않고 환경변수와 마스킹을 사용했다.",
        "AI 생성 코드는 팀원이 실행 결과와 오류를 확인한 후 채택했다.",
        "프롬프트 로그를 조작하지 않고 실제 사용 흐름과 오류 수정 과정을 남기는 것을 원칙으로 했다.",
        "개인 성찰은 AI 작성 금지 항목이므로 본 보고서에서는 작성 지침만 제공한다.",
    ]:
        bullet(doc, item)

    section_title(doc, "8. 결과 및 평가", "[PO2] 구현 결과 검증")
    h2(doc, "8.1 최종 결과물")
    table(
        doc,
        ["산출물", "위치", "설명"],
        [
            ["웹 애플리케이션", "frontend, src/main/java", "React + Spring Boot 기반 통합 스포츠 분석 서비스"],
            ["최종 체크리스트", "docs/FINAL_TEST_CHECKLIST.md", "발표 전 기능별 확인 항목"],
            ["프로젝트 상태 기록", "docs/PROJECT_STATUS_2026-06-01.md", "최종 정리와 검증 결과"],
            ["보고서", "outputs/final-report", "평가기준 기반 최종 보고서"],
            ["발표자료", "outputs/final-report", "12장 PPTX와 발표 대본"],
        ],
        [Cm(3.2), Cm(5), Cm(7.8)],
    )
    h2(doc, "8.2 성능 평가 결과")
    table(
        doc,
        ["평가 항목", "결과", "근거"],
        [
            ["시스템 완성도", "주요 기능 구현 완료", "홈, 경기센터, 상세, 분석, 스포츠 허브, 관리자"],
            ["안정성", "빌드/컴파일 통과", "lint/build/compile PASS"],
            ["보안", "세션 기반 로그인과 관리자 라우트 제한", "ProtectedRoute, AdminRoute, requireAdmin"],
            ["사용성", "시인성, 로고 통일, 불필요 문구 정리", "사용자 피드백 기반 반복 수정"],
            ["한계", "자동화 테스트와 배포 환경은 미완성", "향후 과제로 명시"],
        ],
        [Cm(4), Cm(5), Cm(7)],
    )

    section_title(doc, "9. 성찰 및 학습 결과", "[PO10] 자기주도적 학습")
    callout(
        doc,
        "중요",
        "평가기준에 따르면 성찰 영역은 반드시 학생 본인이 직접 작성해야 하며 AI 사용이 확인되면 성찰 영역 점수가 0점 처리될 수 있다. "
        "따라서 아래 항목은 작성 가이드와 빈 칸으로만 제공한다.",
        "FEE2E2",
    )
    h2(doc, "9.1 기술적 학습 성과")
    paragraph(doc, "[팀원이 직접 작성] 예: React 상태 관리, REST API 연동, JPA/Oracle 매핑, 외부 API 응답 분석에서 배운 점")
    h2(doc, "9.2 AI 협업 역량 성장")
    paragraph(doc, "[팀원이 직접 작성] 예: 초기 단순 프롬프트에서 목표·제약·검증을 포함한 구조화 프롬프트로 발전한 경험")
    h2(doc, "9.3 실패에서 배운 교훈")
    paragraph(doc, "[팀원이 직접 작성] 예: AI가 제안한 코드가 실제 환경에서 실패했을 때 재현 조건을 어떻게 정리했는지")
    h2(doc, "9.4 개인별 성찰")
    table(
        doc,
        ["팀원", "개인 성찰 300~500자"],
        [
            ["조우성", "직접 작성 필요"],
            ["김우림", "직접 작성 필요"],
            ["최수인", "직접 작성 필요"],
        ],
        [Cm(3), Cm(13)],
    )

    section_title(doc, "10. 결론 및 향후 과제", None)
    h2(doc, "10.1 결론")
    paragraph(
        doc,
        "본 프로젝트는 축구, 야구, e스포츠라는 서로 다른 데이터 구조를 하나의 사용자 흐름으로 통합한 웹 서비스이다. "
        "React와 Spring Boot를 분리해 유지보수성을 확보했고, 외부 API와 DB 데이터를 연결해 실제 경기 정보를 제공했다. "
        "AI는 코드 생성 도구가 아니라 설계 검토, 오류 분석, 문서화 보조 도구로 활용했으며, 최종 채택과 검증은 사람이 수행했다.",
    )
    h2(doc, "10.2 향후 과제 및 발전 방향")
    for item in [
        "JUnit/React Testing Library 기반 자동화 테스트 추가",
        "API 실패 시 캐시 데이터 표시와 재시도 정책 개선",
        "모바일 반응형 화면 및 접근성 점검 강화",
        "배포 환경 구성과 운영용 환경변수 관리",
        "AI 분석 결과의 출처와 한계 안내 강화",
    ]:
        bullet(doc, item)

    section_title(doc, "부록 A. 학습성과(PO) 달성도 자가평가", None)
    table(
        doc,
        ["PO", "학습성과", "관련 보고서 섹션", "달성도(1~5)"],
        [
            ["PO1", "컴퓨팅 지식 응용", "5장", "4"],
            ["PO2", "이론/알고리즘 검증", "5장, 8장", "4"],
            ["PO3", "문제 정의 및 모델링", "1장", "4"],
            ["PO4", "최신 도구 활용", "2장, 3장", "4"],
            ["PO5", "현실적 제한조건 고려 설계", "1장, 4장", "4"],
            ["PO6", "팀워크", "6장", "팀원 보완 필요"],
            ["PO7", "의사소통", "3장, 6장, 발표자료", "4"],
            ["PO8", "사회적 영향 이해", "7장, 9장", "팀원 보완 필요"],
            ["PO9", "직업윤리/사회적 책임", "2장, 7장, 9장", "4"],
            ["PO10", "자기주도적 학습", "9장", "직접 작성 필요"],
        ],
        [Cm(2), Cm(6), Cm(5), Cm(3)],
    )

    section_title(doc, "부록 B. 핵심 프롬프트 로그 발췌", None)
    table(
        doc,
        ["구분", "핵심 프롬프트 요지", "관련 문서"],
        [
            ["공통", "React + Spring Boot REST API 전환, 인증/권한, 오류 처리", "docs/COMMON_FEATURES_PROMPT_LOG.md"],
            ["야구", "MLB 일정 동기화, 상세 기록, 라인스코어, AI 리포트", "docs/BASEBALL_PROMPT_LOG.md"],
            ["e스포츠", "Cito API, LCK 오늘 경기, 팀/선수 지표, Groq 분석", "docs/ESPORTS_PROMPT_LOG.md"],
            ["축구", "K리그 데이터 수집, 상세 분석, 순위", "PROMPT_LOG-soccer.md, README-soccer.md"],
        ],
        [Cm(3), Cm(9), Cm(4)],
    )

    section_title(doc, "부록 C. 회의록", None)
    paragraph(doc, "주차별 회의록은 팀원이 실제 회의 날짜와 내용을 기준으로 추가해야 한다.")

    section_title(doc, "부록 D. 참고 자료", None)
    for item in [
        "MLB Stats API",
        "Cito API",
        "TheSportsDB API",
        "Spring Boot Documentation",
        "React Documentation",
        "Groq API Documentation",
        "프로젝트 저장소 내 README.md, docs/FINAL_TEST_CHECKLIST.md, docs/EVALUATION_EVIDENCE_LOG.md",
    ]:
        bullet(doc, item)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
